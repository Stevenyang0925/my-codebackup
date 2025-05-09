"""
Word文档解析器模块
"""

import os
import re
import mammoth
from typing import Dict, Any, List
from .base_parser import BaseParser
from ..utils.logger import get_logger
from ..utils.exceptions import FileParsingError

# 获取 logger 实例
logger = get_logger(__name__)

class WordParser(BaseParser):
    """
    Word文档解析器，用于解析.docx和.doc文件
    """
    
    def get_supported_extensions(self) -> List[str]:
        """
        获取支持的文件扩展名列表
        
        Returns:
            List[str]: 支持的文件扩展名列表
        """
        return ['.docx', '.doc']
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析Word文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict[str, Any]: 解析结果
        """
        if not self.is_supported(file_path):
            logger.error(f"不支持的文件类型传递给 WordParser: {file_path}")
            raise FileParsingError(f"WordParser 不支持的文件类型: {file_path}")
        
        title = self.extract_title(file_path)
        logger.info(f"开始使用 Mammoth 解析 Word 文档: {file_path}")
        
        try:
            # 设置自定义样式映射
            style_map = """
            p[style-name='Heading 1'] => h1:fresh
            p[style-name='Heading 2'] => h2:fresh
            p[style-name='Heading 3'] => h3:fresh
            p[style-name='Heading 4'] => h4:fresh
            p[style-name='标题 1'] => h1:fresh
            p[style-name='标题 2'] => h2:fresh
            p[style-name='标题 3'] => h3:fresh
            p[style-name='标题 4'] => h4:fresh
            p[style-name='Title'] => h1:fresh
            p[style-name='Subtitle'] => h2:fresh
            p[style-name='Normal'] => p:fresh
            p[style-name='List Paragraph'] => p:fresh
            r[style-name='Strong'] => strong
            r[style-name='Emphasis'] => em
            
            # 添加更多样式映射以提高兼容性
            p[outline-level="1"] => h1:fresh
            p[outline-level="2"] => h2:fresh
            p[outline-level="3"] => h3:fresh
            p[outline-level="4"] => h4:fresh
            
            # 处理可能的中文样式名称变体
            p[style-name='一级标题'] => h1:fresh
            p[style-name='二级标题'] => h2:fresh
            p[style-name='三级标题'] => h3:fresh
            p[style-name='四级标题'] => h4:fresh
            """
            
            # 使用mammoth库将Word文档转换为Markdown
            with open(file_path, "rb") as docx_file:
                try:
                    # 使用正确的 API 调用方式
                    result = mammoth.convert_to_markdown(docx_file, style_map=style_map)
                    markdown_content = result.value
                    
                    # 调试：打印 Mammoth 的原始输出
                    logger.debug(f"==== Mammoth 原始输出（前500个字符）====\n{result.value[:500]}\n==== Mammoth 原始输出结束 ====")
                    
                    # 获取转换过程中的警告信息
                    messages = result.messages
                    if messages:
                        logger.warning("Mammoth 转换过程中产生警告:")
                        for message in messages:
                            logger.warning(f"  - {message.message} ({message.type})")
                except Exception as mammoth_error:
                    logger.error(f"Mammoth 转换失败: {mammoth_error}", exc_info=True)
                    raise FileParsingError(f"Mammoth 转换 Word 文件失败: {mammoth_error}") from mammoth_error
            
            # 直接处理markdown内容，不再转换为结构化内容再转回markdown
            logger.info("开始智能处理 Mammoth 生成的 Markdown 内容")
            markdown_content = self._process_markdown_intelligently(markdown_content)
            
            # 返回处理后的markdown内容
            logger.info(f"Word 文档解析和处理完成: {file_path}")
            return {
                "title": title,
                "content": markdown_content,
                "raw_markdown": True  # 标记为原始markdown内容，不需要再转换
            }
        except FileNotFoundError:
            logger.error(f"Word 文件未找到: {file_path}", exc_info=True)
            raise FileParsingError(f"Word 文件未找到: {file_path}")
        except Exception as e:
            # 在实际应用中，应该使用日志记录错误
            logger.exception(f"使用 Mammoth 解析 Word 文档时发生未知错误: {e}")
            # 如果mammoth失败，尝试使用备选方法
            return self._parse_with_python_docx(file_path, title)
    
    def _process_markdown_intelligently(self, markdown_content: str) -> str:
        """
        智能处理markdown内容，基于文本特征识别标题
        
        Args:
            markdown_content: 原始markdown内容
            
        Returns:
            str: 处理后的markdown内容
        """
        logger.debug(f"==== 开始智能处理 Markdown 内容 ====\n原始内容前200个字符: {markdown_content[:200]}")
        
        # 1. 规范化专业术语
        terms = {
            r'\bppg\b': 'PPG',
            r'\badc\b': 'ADC',
            r'\bimu\b': 'IMU',
            r'\bntc\b': 'NTC',
            r'\bble\b': 'BLE',
            r'\bsnr\b': 'SNR',
            r'\bled\b': 'LED',
            r'\baod\b': 'AOD',
            r'\batm\b': 'ATM'
        }
        
        for pattern, replacement in terms.items():
            markdown_content = re.sub(pattern, replacement, markdown_content, flags=re.IGNORECASE)
        
        # 2. 处理标题格式
        lines = markdown_content.split('\n')
        processed_lines = []
        
        logger.debug(f"智能处理：总行数: {len(lines)}")
        
        # 跟踪文档结构
        in_list = False
        prev_line_empty = True  # 文档开始被视为空行之后
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # 跳过空行
            if not line_stripped:
                processed_lines.append(line)
                prev_line_empty = True
                in_list = False  # 空行结束列表
                continue
            
            # 处理已有的标题行，确保格式正确
            if line.startswith('#'):
                logger.debug(f"智能处理：发现已有标题行: {line}")
                # 确保标题格式正确（# 后有空格）
                if not re.match(r'^#+\s', line):
                    line = re.sub(r'^(#+)', r'\1 ', line)
                    logger.debug(f"智能处理：修正标题格式 -> {line}")
                processed_lines.append(line)
                prev_line_empty = False
                in_list = False
                continue
            
            # 检测列表项
            if re.match(r'^[\*\-\+]\s', line_stripped) or re.match(r'^\d+\.\s', line_stripped):
                logger.debug(f"智能处理：发现列表项: {line}")
                in_list = True
                processed_lines.append(line)
                prev_line_empty = False
                continue
            
            # 特殊处理：已知的章节标题
            known_sections = ["息屏显", "冷钱包", "5ATM防水保护", "3D弧面玻璃", "快充加持", "多功能配件", "数字身份"]
            if line_stripped in known_sections:
                logger.debug(f"智能处理：发现已知章节标题: {line} -> ## {line_stripped}")
                line = f"## {line_stripped}"
                processed_lines.append(line)
                prev_line_empty = False
                in_list = False
                continue
            
            # 智能识别可能的标题
            # 标题特征：
            # 1. 前面是空行或文档开始
            # 2. 长度较短（通常不超过50个字符）
            # 3. 不以标点符号结尾（除了冒号）
            # 4. 后面是空行或文档结束
            # 5. 不是列表的一部分
            is_potential_heading = (
                prev_line_empty and  # 前面是空行
                not in_list and  # 不是列表的一部分
                len(line_stripped) <= 50 and  # 长度较短
                not re.search(r'[.,:;?!]$', line_stripped.rstrip(':')) and  # 不以标点符号结尾（除了冒号）
                (i == len(lines) - 1 or not lines[i+1].strip())  # 后面是空行或文档结束
            )
            
            # 检查是否有其他标题特征
            if not is_potential_heading and prev_line_empty and not in_list:
                # 检查是否是孤立的短行（可能是标题）
                if (len(line_stripped) <= 30 and  # 非常短的行
                    not re.search(r'[.。，,;；]', line_stripped) and  # 不以标点符号结尾
                    not re.match(r'^[\*\-\+]\s', line_stripped) and  # 不是列表项
                    not re.match(r'^\d+\.\s', line_stripped)):  # 不是有序列表项
                    is_potential_heading = True
            
            if is_potential_heading:
                logger.debug(f"智能处理：发现潜在标题: {line} -> ## {line_stripped}")
                # 判断标题级别：如果这是文档的第一个潜在标题，且前面没有 # 标题，则可能是一级标题
                is_first_heading = all(not l.startswith('#') for l in lines[:i])
                
                if is_first_heading and i < 5:  # 如果是文档前几行的第一个标题
                    line = f"# {line_stripped}"
                else:  # 否则视为二级标题
                    line = f"## {line_stripped}"
                
                processed_lines.append(line)
                prev_line_empty = False
                continue
            
            # 处理其他行
            processed_lines.append(line)
            prev_line_empty = False
        
        # 3. 重新组合内容
        markdown_content = '\n'.join(processed_lines)
        
        # 4. 确保段落之间有空行
        markdown_content = re.sub(r'([^\n])\n([^#\n\-\*\+\d])', r'\1\n\n\2', markdown_content)
        
        # 5. 确保标题前后有空行
        markdown_content = re.sub(r'([^\n])\n(#+\s)', r'\1\n\n\2', markdown_content)
        markdown_content = re.sub(r'(#+\s[^\n]+)\n([^#\n\-\*\+\d])', r'\1\n\n\2', markdown_content)
        
        logger.debug(f"处理后内容前200个字符: {markdown_content[:200]}\n==== 智能处理完成 ====")
        
        return markdown_content
    
    def _post_process_markdown(self, markdown_content: str) -> str:
        """
        对生成的markdown内容进行后处理
        
        Args:
            markdown_content: 原始markdown内容
            
        Returns:
            str: 处理后的markdown内容
        """
        logger.debug("开始 Markdown 后处理")
        # 规范化专业术语
        terms = {
            r'\bppg\b': 'PPG',
            r'\badc\b': 'ADC',
            r'\bimu\b': 'IMU',
            r'\bntc\b': 'NTC',
            r'\bble\b': 'BLE',
            r'\bsnr\b': 'SNR',
            r'\bled\b': 'LED'
        }
        
        for pattern, replacement in terms.items():
            markdown_content = re.sub(pattern, replacement, markdown_content, flags=re.IGNORECASE)
        
        # 修复重复标题
        lines = markdown_content.split('\n')
        unique_lines = []
        seen_headings = set()
        duplicates_removed = 0
        
        for line in lines:
            if line.startswith('#'):
                heading_text = line.strip()
                if heading_text in seen_headings:
                    duplicates_removed += 1
                    continue
                seen_headings.add(heading_text)
            unique_lines.append(line)
        
        if duplicates_removed > 0:
            logger.debug(f"后处理：移除了 {duplicates_removed} 个重复标题")
        
        # 修复动态采集策略部分的格式
        markdown_content = '\n'.join(unique_lines)
        markdown_content = self._fix_collection_strategy_format_md(markdown_content)
        
        # 修复Mammoth转换中的常见问题
        markdown_content = self._fix_mammoth_conversion_issues(markdown_content)
        
        # 确保段落之间有空行
        markdown_content = re.sub(r'([^\n])\n([^#\n-])', r'\1\n\n\2', markdown_content)
        
        # 确保标题格式正确（标题前后有空行）
        markdown_content = re.sub(r'([^\n])\n(#+\s)', r'\1\n\n\2', markdown_content)
        markdown_content = re.sub(r'(#+\s[^\n]+)\n([^#\n])', r'\1\n\n\2', markdown_content)
        
        logger.debug("Markdown 后处理完成")
        return markdown_content.strip()
    
    def _fix_mammoth_conversion_issues(self, markdown_content: str) -> str:
        """
        修复Mammoth转换中的常见问题
        
        Args:
            markdown_content: mammoth转换的markdown内容
            
        Returns:
            str: 修复后的markdown内容
        """
        logger.debug("开始修复 Mammoth 转换问题")
        # 1. 修复重复的空行
        markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
        
        # 2. 去除不必要的转义字符
        markdown_content = re.sub(r'\\([.#\-*_])', r'\1', markdown_content)
        
        # 3. 修复列表项格式问题 - 确保列表项有正确的格式
        markdown_content = re.sub(r'\n\*\s*([^\n]+)', r'\n* \1', markdown_content)
        
        # 4. 修复标题格式问题 - 确保#后有空格
        markdown_content = re.sub(r'#([^#\s])', r'# \1', markdown_content)
        
        # 5. 识别数字列表开头的标题，转换为标题格式
        lines = markdown_content.split('\n')
        new_lines = []
        current_section = None
        section_level = 1
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 判断是否是标题行
            title_match = re.match(r'^(\d+\.\s*)__(.+?)__$', line)
            if title_match and len(line) < 100:  # 避免将长文本误识别为标题
                number = title_match.group(1)
                title = title_match.group(2)
                # 将其转换为Markdown标题
                new_lines.append(f"\n## {title}")
                current_section = title
                section_level = 2
                i += 1
                continue
            
            # 判断是否是二级标题行
            subtitle_match = re.match(r'^(\d+\.\s*)(.+?):$', line)
            if subtitle_match and len(line) < 80 and not re.search(r'[.。，,;；]', subtitle_match.group(2)):
                number = subtitle_match.group(1)
                title = subtitle_match.group(2)
                # 如果标题包含"模式"，则视为三级标题
                if "模式" in title:
                    new_lines.append(f"\n### {title}")
                    section_level = 3
                else:
                    new_lines.append(f"\n### {title}")
                    section_level = 3
                i += 1
                continue
            
            # 识别列表项
            list_match = re.match(r'^(\d+\.\s*)(.+)$', line)
            if list_match and len(line) < 200:  # 一般列表项不会太长
                content = list_match.group(2)
                
                # 如果内容看起来像是属性描述，则转换为无序列表项
                if ":" in content or "-" in content or content.startswith(("灵敏度", "型号", "范围", "分辨率", "精度", "功耗", "防水", "蓝牙")):
                    # 将描述格式化为无序列表项
                    formatted_content = content.replace(":", ": ")
                    new_lines.append(f"- {formatted_content}")
                else:
                    # 保持原始格式
                    new_lines.append(line)
                i += 1
                continue
            
            # 处理特殊情况：已处理为列表但缺少"- "前缀的行
            if (current_section and 
                line and 
                not line.startswith(("#", "-", "*", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "0.")) and 
                ":" in line and
                len(line) < 100):
                # 可能是列表项但缺少标记
                new_lines.append(f"- {line}")
                i += 1
                continue
            
            # 对于动态采集策略部分的特殊处理
            if current_section and "动态采集策略" in current_section:
                # 如果是模式描述行但没有被识别为标题
                mode_match = re.match(r'^(.*?)模式\((.*?)\)：?$', line)
                if mode_match:
                    mode_name = mode_match.group(1).strip()
                    mode_time = mode_match.group(2).strip()
                    new_lines.append(f"\n### {mode_name}模式({mode_time})")
                    i += 1
                    continue
                
                # 处理模式下的项目
                item_match = re.match(r'^-?\s*([A-Za-z]+):\s*(.+)$', line)
                if item_match:
                    item_name = item_match.group(1)
                    item_desc = item_match.group(2)
                    # 确保使用无序列表格式
                    new_lines.append(f"- {item_name}: {item_desc}")
                    i += 1
                    continue
            
            # 默认情况：保留原始行
            new_lines.append(line)
            i += 1
        
        # 重新组合内容
        processed_content = '\n'.join(new_lines)
        
        # 6. 修复表格格式
        table_rows = re.findall(r'(\|[^\n]+\|)', processed_content)
        if table_rows:
            for row in table_rows:
                # 修复单元格之间的间距问题
                fixed_row = re.sub(r'\|\s*([^|]*?)\s*\|', r'| \1 |', row)
                processed_content = processed_content.replace(row, fixed_row)
            
            # 如果存在表格但缺少分隔行，添加分隔行
            table_start_indices = [m.start() for m in re.finditer(r'\|[^\n]+\|\n', processed_content)]
            for i in table_start_indices:
                end_of_line = processed_content.find('\n', i)
                if end_of_line > 0:
                    first_row = processed_content[i:end_of_line]
                    cells_count = first_row.count('|') - 1
                    if cells_count > 0:
                        separator_row = '|' + '|'.join([' --- ' for _ in range(cells_count)]) + '|\n'
                        next_char_pos = end_of_line + 1
                        if next_char_pos < len(processed_content) and not processed_content[next_char_pos:].startswith('|'):
                            processed_content = processed_content[:end_of_line+1] + separator_row + processed_content[end_of_line+1:]
        
        # 7. 最后进行一些清理和优化
        # 移除多余的空行
        processed_content = re.sub(r'\n{3,}', '\n\n', processed_content)
        # 移除行首空格
        processed_content = re.sub(r'(?m)^\s+', '', processed_content)
        
        # 确保标题格式正确
        processed_content = re.sub(r'(?m)^(#+)([^ ])', r'\1 \2', processed_content)
        
        # 确保列表项格式正确
        processed_content = re.sub(r'(?m)^(-|\*)([^ ])', r'\1 \2', processed_content)
        
        # 确保段落之间有空行
        processed_content = re.sub(r'([^\n])\n([^#\n-])', r'\1\n\n\2', processed_content)
        
        logger.debug("修复 Mammoth 转换问题完成")
        return processed_content.strip()
    
    def _fix_collection_strategy_format_md(self, markdown_content: str) -> str:
        """
        修复动态采集策略部分的格式
        
        Args:
            markdown_content: markdown内容
            
        Returns:
            str: 修复后的markdown内容
        """
        logger.debug("开始修复动态采集策略格式")
        # 查找动态采集策略部分
        strategy_pattern = r'(#+\s*动态采集策略.*?)(?=#+\s|$)'
        strategy_match = re.search(strategy_pattern, markdown_content, re.DOTALL)
        
        if not strategy_match:
            return markdown_content
        
        strategy_section = strategy_match.group(1)
        
        # 修复模式格式
        fixed_section = re.sub(
            r'(\d+\.\s*)(.*?)模式\((.*?)\):(.*?)(?=\d+\.|$)', 
            r'\1\2模式(\3):\n\4', 
            strategy_section, 
            flags=re.DOTALL
        )
        
        # 修复列表项格式
        fixed_section = re.sub(
            r'([^-])(PPG|温度|IMU):', 
            r'\1- \2:', 
            fixed_section
        )
        
        # 替换原始部分
        return markdown_content.replace(strategy_section, fixed_section)
    
    def _markdown_to_structured_content(self, markdown_content: str) -> List[Dict[str, Any]]:
        """
        将markdown内容转换为结构化内容
        
        Args:
            markdown_content: markdown内容
            
        Returns:
            List[Dict[str, Any]]: 结构化内容
        """
        logger.debug("开始将 Markdown 转换为结构化内容")
        content = []
        lines = markdown_content.split('\n')
        
        i = 0
        current_list_items = []
        current_list_type = None
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                # 如果当前有列表项，并且遇到空行，说明列表结束
                if current_list_items:
                    content.append({
                        "type": "list",
                        "content": "",
                        "additional_info": current_list_items,
                        "list_type": current_list_type or "unordered"
                    })
                    current_list_items = []
                    current_list_type = None
                i += 1
                continue
            
            # 处理标题 (# 标题)
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                heading_level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                
                # 结束之前的列表
                if current_list_items:
                    content.append({
                        "type": "list",
                        "content": "",
                        "additional_info": current_list_items,
                        "list_type": current_list_type or "unordered"
                    })
                    current_list_items = []
                    current_list_type = None
                
                content.append({
                    "type": "heading",
                    "content": heading_text,
                    "level": heading_level
                })
                
                i += 1
                continue
            
            # 处理无序列表 (- 项目)
            unordered_list_match = re.match(r'^[-*]\s+(.+)$', line)
            if unordered_list_match:
                item_text = unordered_list_match.group(1).strip()
                
                # 如果之前是有序列表，先结束之前的列表
                if current_list_type == "ordered" and current_list_items:
                    content.append({
                        "type": "list",
                        "content": "",
                        "additional_info": current_list_items,
                        "list_type": "ordered"
                    })
                    current_list_items = []
                
                current_list_type = "unordered"
                current_list_items.append(item_text)
                i += 1
                continue
            
            # 处理有序列表 (1. 项目)
            ordered_list_match = re.match(r'^(\d+)\.?\s+(.+)$', line)
            if ordered_list_match:
                item_text = ordered_list_match.group(2).strip()
                
                # 如果之前是无序列表，先结束之前的列表
                if current_list_type == "unordered" and current_list_items:
                    content.append({
                        "type": "list",
                        "content": "",
                        "additional_info": current_list_items,
                        "list_type": "unordered"
                    })
                    current_list_items = []
                
                current_list_type = "ordered"
                current_list_items.append(item_text)
                i += 1
                continue
            
            # 处理表格 (| 单元格 | 单元格 |)
            if line.startswith('|') and line.endswith('|'):
                # 结束之前的列表
                if current_list_items:
                    content.append({
                        "type": "list",
                        "content": "",
                        "additional_info": current_list_items,
                        "list_type": current_list_type or "unordered"
                    })
                    current_list_items = []
                    current_list_type = None
                
                # 收集表格行
                table_data = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_row = lines[i].strip()
                    # 跳过分隔行 (| --- | --- |)
                    if not re.search(r'\|[\s\-:]+\|', table_row):
                        cells = [cell.strip() for cell in table_row.strip('|').split('|')]
                        table_data.append(cells)
                    i += 1
                    
                    # 跳过表格后的空行
                    while i < len(lines) and not lines[i].strip():
                        i += 1
                
                if table_data:
                    content.append({
                        "type": "table",
                        "content": "",
                        "additional_info": table_data
                    })
                
                continue
            
            # 处理其他内容（默认为普通文本）
            # 结束之前的列表
            if current_list_items:
                content.append({
                    "type": "list",
                    "content": "",
                    "additional_info": current_list_items,
                    "list_type": current_list_type or "unordered"
                })
                current_list_items = []
                current_list_type = None
            
            # 如果看起来是代码块
            if line.startswith('```'):
                code_content = []
                language = line[3:].strip()
                i += 1  # 跳过```行
                
                while i < len(lines) and not lines[i].strip() == '```':
                    code_content.append(lines[i])
                    i += 1
                
                if i < len(lines):  # 跳过结束的```
                    i += 1
                    
                content.append({
                    "type": "code",
                    "content": "\n".join(code_content),
                    "additional_info": language if language else "plaintext"
                })
            else:
                # 普通文本
                content.append({
                    "type": "text",
                    "content": line
                })
                i += 1
        
        # 处理末尾可能的列表
        if current_list_items:
            content.append({
                "type": "list",
                "content": "",
                "additional_info": current_list_items,
                "list_type": current_list_type or "unordered"
            })
        
        logger.debug("Markdown 转换为结构化内容完成")
        return content
    
    def _parse_with_python_docx(self, file_path: str, title: str) -> Dict[str, Any]:
        """
        使用python-docx库解析Word文档（原始实现）
        
        Args:
            file_path: 文件路径
            title: 文档标题
            
        Returns:
            Dict[str, Any]: 解析结果
        """
        logger.warning(f"尝试使用 python-docx 解析: {file_path}")
        content = []
        
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 提取标题（使用文档的第一个段落作为标题）
            if doc.paragraphs and doc.paragraphs[0].text:
                title = doc.paragraphs[0].text
            
            # 增强标题识别和层级结构处理
            current_section = None
            current_subsection = None
            list_items = []
            in_list = False
            
            # 跟踪已处理的标题，避免重复
            processed_headings = set()
            
            # 标点符号映射表，用于统一标点符号
            punctuation_map = {
                '：': ':',
                '，': ',',
                '。': '.',
                '（': '(',
                '）': ')',
                '、': ',',
            }
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    # 空段落可能表示列表结束
                    if in_list and list_items:
                        content.append({
                            "type": "list",
                            "content": "",
                            "additional_info": list_items,
                            "list_type": "unordered"
                        })
                        list_items = []
                        in_list = False
                    continue
                
                # 初始化变量
                is_heading = False
                heading_level = 0
                is_list_item = False  # 确保变量在每次循环开始时初始化
                
                # 标准化标点符号
                for cn_punct, en_punct in punctuation_map.items():
                    text = text.replace(cn_punct, en_punct)
                
                # 增强标题识别 - 检查样式和其他特征
                # 改进标题检测逻辑
                is_heading = False
                heading_level = 0
                
                # 1. 防止重复标题
                if text.lower() == title.lower() and len(processed_headings) > 0:
                    continue  # 跳过与文档标题重复的标题
                
                # 2. 改进标题级别判断
                # 改进标题识别逻辑
                # 修复para.style.name可能为None的问题
                if para.style and para.style.name and (para.style.name.startswith(('Heading', '标题')) or para.style.name in ['Title', 'Subtitle']):
                    try:
                        heading_level = min(int(re.sub(r'[^\d]', '', para.style.name)) or 1, 3)
                        is_heading = True
                    except ValueError:
                        heading_level = 1
                        is_heading = True
                elif len(text) < 50 and not re.search(r'[.,:;?!]$', text):
                    ...
                
                # 改进列表项识别
                list_prefix = re.match(r'^([-•*·]|\d+\.)\s+', text)
                is_list_item = False  # 确保初始化
                
                if list_prefix:
                    is_list_item = True
                    item_text = text[list_prefix.end():]
                    # 处理带时间范围的列表项
                    if re.search(r'\d{2}:\d{2}-\d{2}:\d{2}', text):
                        item_text = text
                elif ':' in text:
                    parts = re.split(r':', text, 1)
                    if len(parts) == 2 and parts[0].strip() and parts[1].strip():
                        is_list_item = True
                        item_text = f"{parts[0].strip()}: {parts[1].strip()}"
                
                # 处理列表项
                if is_list_item:
                    list_items.append(item_text)
                    in_list = True
                    continue
                
                # 如果之前在处理列表，先结束列表
                if in_list and list_items:
                    content.append({
                        "type": "list",
                        "content": "",
                        "additional_info": list_items,
                        "list_type": "unordered"
                    })
                    list_items = []
                    in_list = False
                
                # 处理普通段落
                content.append({
                    "type": "text",
                    "content": text
                })
            
            # 处理文档末尾可能的未完成列表
            if in_list and list_items:
                content.append({
                    "type": "list",
                    "content": "",
                    "additional_info": list_items,
                    "list_type": "unordered"
                })
            
            # 处理表格
            for table in doc.tables:
                table_data = []
                
                # 提取表格数据
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    content.append({
                        "type": "table",
                        "content": "",
                        "additional_info": table_data
                    })
            
            # 后处理 - 规范化专业术语大小写
            self._normalize_terminology(content)
            
            # 后处理 - 修复动态采集策略部分的格式
            self._fix_collection_strategy_format(content)
            
            # 在parse方法的return之前添加：
            content = self._fix_format_issues(content)
            
            logger.info(f"使用 python-docx 解析成功: {file_path}")
            return {
                "title": title,
                "content": content
            }
        except ImportError:
            logger.error("python-docx 未安装，无法使用备选解析方案")
            raise FileParsingError("python-docx 未安装，无法解析 Word 文件")
        except Exception as e:
            logger.exception(f"使用 python-docx 解析时发生错误: {e}")
            raise FileParsingError(f"使用 python-docx 解析时发生错误: {e}") from e
    
    def _fix_collection_strategy_format(self, content):
        """
        修复动态采集策略部分的格式问题
        
        Args:
            content: 内容列表
        """
        # 查找动态采集策略部分
        strategy_index = -1
        for i, item in enumerate(content):
            if item.get("type") == "heading" and "动态采集策略" in item.get("content", ""):
                strategy_index = i
                break
        
        if strategy_index == -1:
            return
        
        # 查找策略部分的列表
        for i in range(strategy_index + 1, len(content)):
            if content[i].get("type") == "list" and content[i].get("additional_info"):
                list_items = content[i].get("additional_info", [])
                new_list_items = []
                
                # 处理复杂的列表项，将其分解为子列表
                for item in list_items:
                    if "模式" in item and ":" in item:
                        # 这是一个模式标题
                        mode_parts = item.split(":", 1)
                        mode_name = mode_parts[0].strip()
                        
                        # 添加模式标题作为单独的列表项
                        new_list_items.append(mode_name + ":")
                        
                        # 处理模式详情
                        if len(mode_parts) > 1:
                            details = mode_parts[1].strip()
                            # 分割详情中的不同部分
                            detail_parts = re.split(r'\s*-\s*', details)
                            for part in detail_parts:
                                if part.strip():
                                    new_list_items.append("  - " + part.strip())
                    else:
                        new_list_items.append(item)
                
                # 更新列表项
                content[i]["additional_info"] = new_list_items
    
    def _normalize_terminology(self, content):
        """
        规范化专业术语的大小写
        
        Args:
            content: 内容列表
        """
        # 扩展术语词典
        terms = {
            "ppg": "PPG",
            "adc": "ADC",
            "imu": "IMU",
            "ntc": "NTC",
            "ble": "BLE",
            "snr": "SNR",
            "led": "LED"
        }
        
        for item in content:
            if "content" in item and isinstance(item["content"], str):
                for term_lower, term_upper in terms.items():
                    # 使用正则表达式确保只替换独立的词，而不是词的一部分
                    item["content"] = re.sub(
                        r'\b' + term_lower + r'\b', 
                        term_upper, 
                        item["content"], 
                        flags=re.IGNORECASE
                    )
            
            if "additional_info" in item and isinstance(item["additional_info"], list):
                for i, info in enumerate(item["additional_info"]):
                    if isinstance(info, str):
                        for term_lower, term_upper in terms.items():
                            item["additional_info"][i] = re.sub(
                                r'\b' + term_lower + r'\b', 
                                term_upper, 
                                info, 
                                flags=re.IGNORECASE
                            )

    def _fix_format_issues(self, content):
        """修复常见的格式问题"""
        new_content = []
        last_was_heading = False
        
        for item in content:
            # 跳过空内容
            if not item.get("content") and not item.get("additional_info"):
                continue
                
            # 修复重复标题
            if item.get("type") == "heading":
                if last_was_heading and item["level"] == 1:
                    continue  # 跳过连续的一级标题
                last_was_heading = True
            else:
                last_was_heading = False
                
            # 修复动态采集策略的列表格式
            if (item.get("type") == "list" and 
                any('模式' in s for s in item.get("additional_info", []))):
                item["list_type"] = "ordered"
                
            new_content.append(item)
        
        return new_content

    def _is_duplicate_heading(self, text, processed_headings):
        """改进的标题去重检查"""
        text_lower = text.lower().strip(' :：')
        for heading in processed_headings:
            if text_lower == heading.lower().strip(' :：'):
                return True
        return False

    def extract_title(self, file_path: str) -> str:
        """
        从Word文档中提取标题
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文档标题
        """
        try:
            # 尝试从文档内容提取标题
            import docx
            try:
                doc = docx.Document(file_path)
                # 使用第一个段落作为标题，如果是标题样式
                if doc.paragraphs and doc.paragraphs[0].text.strip():
                    first_para = doc.paragraphs[0]
                    if first_para.style and first_para.style.name and (first_para.style.name.startswith(('Heading', '标题')) or first_para.style.name in ['Title', 'Subtitle']):
                        return first_para.text.strip()
                        
                # 尝试使用文档属性中的标题
                if doc.core_properties.title:
                    return doc.core_properties.title
            except:
                pass
        except:
            pass
        
        # 如果无法从内容提取标题，使用文件名
        import os
        base_name = os.path.basename(file_path)
        title, _ = os.path.splitext(base_name)
        return title

    def test_conversion(self, file_path: str, output_path: str | None = None) -> str | None:
        """
        测试Word文档转换效果，将结果保存为Markdown文件
        
        Args:
            file_path: Word文档路径
            output_path: 输出路径，如果为None则使用原文件名加上时间戳
            
        Returns:
            str: 保存的Markdown文件路径
        """
        import time
        from datetime import datetime
        
        try:
            # 设置自定义样式映射
            style_map = """
            p[style-name='Heading 1'] => h1:fresh
            p[style-name='Heading 2'] => h2:fresh
            p[style-name='Heading 3'] => h3:fresh
            p[style-name='Heading 4'] => h4:fresh
            p[style-name='标题 1'] => h1:fresh
            p[style-name='标题 2'] => h2:fresh
            p[style-name='标题 3'] => h3:fresh
            p[style-name='标题 4'] => h4:fresh
            p[style-name='Title'] => h1:fresh
            p[style-name='Subtitle'] => h2:fresh
            p[style-name='Normal'] => p:fresh
            p[style-name='List Paragraph'] => p:fresh
            r[style-name='Strong'] => strong
            r[style-name='Emphasis'] => em
            
            # 添加更多样式映射以提高兼容性
            p[outline-level="1"] => h1:fresh
            p[outline-level="2"] => h2:fresh
            p[outline-level="3"] => h3:fresh
            p[outline-level="4"] => h4:fresh
            
            # 处理可能的中文样式名称变体
            p[style-name='一级标题'] => h1:fresh
            p[style-name='二级标题'] => h2:fresh
            p[style-name='三级标题'] => h3:fresh
            p[style-name='四级标题'] => h4:fresh
            """
            
            # 使用mammoth库将Word文档转换为Markdown
            with open(file_path, "rb") as docx_file:
                print(f"正在转换文件：{file_path}")
                result = mammoth.convert_to_markdown(
                    docx_file,
                    style_map=style_map
                )
                markdown_content = result.value
                
                # 调试：打印 Mammoth 的原始输出
                print("\n==== Mammoth 原始输出（前500个字符）====")
                print(result.value[:500])
                print("==== Mammoth 原始输出结束 ====\n")
                
                # 获取转换过程中的警告信息
                messages = result.messages
                if messages:
                    print("\n警告信息:")
                    for message in messages:
                        print(f"- {message.message} ({message.type})")
                
                # 后处理markdown内容
                print("\n开始后处理Markdown内容...")
                processed_content = self._post_process_markdown(markdown_content)
                
                # 如果没有指定输出路径，则使用原文件名加上时间戳
                if output_path is None:
                    filename = os.path.basename(file_path)
                    base_name, _ = os.path.splitext(filename)
                    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                    output_path = f"{base_name}_{timestamp}.md"
                
                # 保存处理后的Markdown内容
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(processed_content)
                    
                print(f"\n转换完成，结果已保存到：{output_path}")
                print(f"Markdown内容预览 (前300个字符):")
                print("-" * 50)
                print(processed_content[:300] + "..." if len(processed_content) > 300 else processed_content)
                print("-" * 50)
                
                return output_path
        except Exception as e:
            print(f"转换过程中出错: {e}")
            import traceback
            traceback.print_exc()
            return None